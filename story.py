import json
from cfonts import render
from rich.console import Console
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

passages = {}
readerpath = []
console = Console()

doc = Document() # Creates a new word document.

try:
    with open('passages.json') as p: # Loads the passages file into dictonary object.
        passages = json.load(p)
except FileNotFoundError:
    console.print("Could not find the passages file.  Please make sure it exists!", style="bold red")
    exit(0)
    
def serialize(path):
    fileText = ""
    for option in path:
        if(option == "end"):
            break
        
        fileText += "\n\n" + passages[0][option]["passage"]
    
    return fileText

def writeStory(fileText, filePath):
    try:
        doc.save("story.docx")
    except IOError:
        console.print("There was a problem writing the file.", style="bold red")        
        
def main():
    messageoutput = render('Create A Story!', colors=['red', 'yellow'], align='center')
    print(messageoutput)
    print()
    
    title = doc.add_heading('Create A Story!', level=1) # Creates a heading in the word document
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Loops through story until an end result is reached.
    current_passage = "begin" # The tag for the beginning passage
    while current_passage != "end": # The tag for the end of story
        current_passage = do_story(current_passage)
        
    writeStory(serialize(readerpath), "story.txt")
    
def do_story(current_passage):
    # Gets the passage from the global passages list, displays the passage, then returns the next selected passage from the user.
    passage = passages[0][current_passage]
    readerpath.append(current_passage)
    print_passage(passage["passage"])
    doc.add_paragraph(passage["passage"])
    return get_option(passage["options"])

def print_passage(passage):
    # Prints the given passage
    console.print(passage, style="blue")
    print()
    
def get_option(options):
    # Iterate through list of options, enumerate them, 
    # then ask for the selection from the user.  
    # Each option is the next passage to go to.
    i = 0
    optionsavail = []
    for option in options: # Enumerate and print each option
        i = i + 1
        console.print(f"{i}: {options[option]}", style="bold blue")
        optionsavail.append(option)
    
    optionnumber = get_option_number(len(optionsavail))
    
    return optionsavail[optionnumber - 1] # Return the option name

def get_option_number(max_option):
    # Gets the option number from the user bounded by max_option and 0
    optionnumber = 0
    while optionnumber < 1 or optionnumber > max_option:
        optionnumber = int(input("Please select a path: "))
        if optionnumber < 1 or optionnumber > max_option: # If the option selected is outside of bounds, loop again and display error
            print("Invalid Option!")
        
    return optionnumber

    
##############
# Test Suite #
##############
# test_path = ["first", "second", "fourth", "end"] # Tests supposed options by the user.

# text = serialize(test_path)
# writeStory(text, "teststory.txt")

if __name__ == '__main__':
    main()
